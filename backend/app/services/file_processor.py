import logging
import os
import shutil
import pdfplumber
from PIL import Image

logger = logging.getLogger(__name__)

# Check if Tesseract is available at import time
_tesseract_cmd = shutil.which("tesseract")
if not _tesseract_cmd:
    # Check common Windows install path
    _win_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.isfile(_win_path):
        _tesseract_cmd = _win_path

_tesseract_available = _tesseract_cmd is not None
if _tesseract_available:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = _tesseract_cmd
    logger.info(f"Tesseract found at: {_tesseract_cmd}")
else:
    logger.warning("Tesseract not found — image OCR will be unavailable")


def extract_text(file_path: str, file_type: str) -> str:
    if file_type in ("pdf", "application/pdf"):
        return _extract_pdf(file_path)
    elif file_type in ("png", "jpg", "jpeg", "image/png", "image/jpeg", "image/jpg"):
        return _extract_image_ocr(file_path)
    elif file_type in ("doc", "docx"):
        return _extract_word(file_path)
    elif file_type in ("xls", "xlsx"):
        return _extract_excel(file_path)
    elif file_type in ("txt", "text/plain"):
        return _read_text(file_path)
    else:
        return _read_text(file_path)


def _extract_pdf(pdf_path: str) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    # If no text extracted (scanned/image PDF), fall back to OCR via PyMuPDF + Tesseract
    if not text.strip() and _tesseract_available:
        logger.info(f"PDF has no selectable text, falling back to OCR: {pdf_path}")
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(pdf_path)
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                if page_text:
                    text += page_text + "\n"
            doc.close()
            logger.info(f"OCR extracted {len(text)} chars from {pdf_path}")
        except ImportError:
            logger.warning("PyMuPDF not installed — cannot OCR scanned PDFs. Install with: pip install pymupdf")
        except Exception as e:
            logger.error(f"OCR fallback failed for PDF: {e}")

    return text.strip()


def _extract_image_ocr(image_path: str) -> str:
    if not _tesseract_available:
        return "[OCR unavailable — Tesseract is not installed. Please install Tesseract-OCR and add it to PATH.]"
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    return text.strip()


def _extract_word(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def _extract_excel(file_path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"--- Sheet: {sheet} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells).strip()
            if line:
                parts.append(line)
    wb.close()
    return "\n".join(parts).strip()


def _read_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()
